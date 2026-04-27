"""Word executive memo generator — McKinsey-styled .docx output.

Produces a structured memo with:
- Title page with project name and date
- Executive summary (SCR framework)
- Key recommendation box
- Supporting evidence sections (from MECE branches)
- Next steps / action items
- Sources appendix

Uses calibrated McKinsey styling: Arial, #002960 navy, clean layout.
"""
from __future__ import annotations
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# McKinsey color palette (calibrated from 37 real PDFs)
NAVY = RGBColor(0x00, 0x29, 0x60)        # #002960
BLUE = RGBColor(0x00, 0x65, 0xBD)        # #0065BD
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)   # #333333
MED_GRAY = RGBColor(0x66, 0x66, 0x66)    # #666666
LIGHT_BG = RGBColor(0xF0, 0xF4, 0xF8)    # #F0F4F8 — subtle blue-gray
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _add_styled_para(doc: Document, text: str, style_name: str = "Normal",
                     bold: bool = False, color: RGBColor = None,
                     size: int = None, alignment=None, space_after: int = None):
    """Add a paragraph with optional inline styling."""
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.name = "Arial"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def _add_recommendation_box(doc: Document, recommendation: str):
    """Add a highlighted recommendation box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "002960")

    # Label
    label_para = cell.paragraphs[0]
    label_run = label_para.add_run("KEY RECOMMENDATION")
    label_run.font.name = "Arial"
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(0x80, 0xBB, 0xE0)
    label_para.paragraph_format.space_after = Pt(4)

    # Recommendation text
    rec_para = cell.add_paragraph()
    rec_run = rec_para.add_run(recommendation)
    rec_run.font.name = "Arial"
    rec_run.font.size = Pt(11)
    rec_run.font.color.rgb = WHITE
    rec_run.bold = True
    rec_para.paragraph_format.space_after = Pt(6)

    # Table border styling via XML
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        border = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "0", qn("w:color"): "002960",
        })
        borders.append(border)
    tblPr.append(borders)

    doc.add_paragraph()  # spacing


def _add_scr_section(doc: Document, label: str, text: str, color: RGBColor):
    """Add a Situation/Complication/Resolution paragraph pair."""
    if not text:
        return
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}: ")
    label_run.font.name = "Arial"
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = color

    text_run = para.add_run(text)
    text_run.font.name = "Arial"
    text_run.font.size = Pt(10)
    text_run.font.color.rgb = DARK_GRAY
    para.paragraph_format.space_after = Pt(6)


def generate_executive_memo(data: dict, output_path: str):
    """Generate a McKinsey-styled executive memo as .docx.

    Args:
        data: dict with keys: project, storyline, slides, research_brief
        output_path: path for the .docx file
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    project = data["project"]
    storyline = data.get("storyline")
    slides = data.get("slides", [])
    research_brief = data.get("research_brief")

    # ── Title section ──
    _add_styled_para(doc, "EXECUTIVE MEMO", bold=True, color=MED_GRAY,
                     size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    _add_styled_para(doc, project["name"], bold=True, color=NAVY,
                     size=18, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    date_str = datetime.now().strftime("%B %d, %Y")
    _add_styled_para(doc, f"Prepared: {date_str}", color=MED_GRAY,
                     size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    if project.get("description"):
        _add_styled_para(doc, project["description"], color=MED_GRAY,
                         size=9, space_after=12)

    # Horizontal rule
    _add_styled_para(doc, "─" * 72, color=RGBColor(0xCC, 0xCC, 0xCC), size=6, space_after=12)

    # ── Key Recommendation ──
    if storyline and storyline.get("key_recommendation"):
        _add_recommendation_box(doc, storyline["key_recommendation"])

    # ── Executive Summary (SCR) ──
    if storyline:
        _add_styled_para(doc, "EXECUTIVE SUMMARY", bold=True, color=NAVY, size=12, space_after=8)
        _add_scr_section(doc, "Situation", storyline.get("situation", ""), BLUE)
        _add_scr_section(doc, "Complication", storyline.get("complication", ""), RGBColor(0xC0, 0x39, 0x2B))
        _add_scr_section(doc, "Resolution", storyline.get("resolution", ""), RGBColor(0x27, 0xAE, 0x60))
        doc.add_paragraph()  # spacing

    # ── Supporting Analysis (from slides) ──
    content_slides = [s for s in slides if s.get("slide_type") in (
        "content_text", "content_chart", "content_hybrid", "recommendation",
    )]

    if content_slides:
        _add_styled_para(doc, "SUPPORTING ANALYSIS", bold=True, color=NAVY, size=12, space_after=8)

        for slide in content_slides:
            title = slide.get("action_title", "")
            if title:
                _add_styled_para(doc, title, bold=True, color=DARK_GRAY, size=10, space_after=4)

            bullets = slide.get("bullets", [])
            for bullet in bullets:
                if isinstance(bullet, dict):
                    prefix = bullet.get("bold_prefix", "")
                    text = bullet.get("text", "")
                    para = doc.add_paragraph(style="List Bullet")
                    if prefix:
                        bold_run = para.add_run(f"{prefix} ")
                        bold_run.font.name = "Arial"
                        bold_run.bold = True
                        bold_run.font.size = Pt(9)
                    text_run = para.add_run(text)
                    text_run.font.name = "Arial"
                    text_run.font.size = Pt(9)
                elif isinstance(bullet, str):
                    para = doc.add_paragraph(bullet, style="List Bullet")
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

            # Chart annotation
            chart = slide.get("chart")
            if chart and isinstance(chart, dict) and chart.get("so_what"):
                _add_styled_para(doc, f"Key insight: {chart['so_what']}", color=BLUE, size=9, space_after=4)

            doc.add_paragraph()  # spacing between sections

    # ── Next Steps (from next_steps slides) ──
    next_steps_slides = [s for s in slides if s.get("slide_type") == "next_steps"]
    if next_steps_slides:
        _add_styled_para(doc, "NEXT STEPS", bold=True, color=NAVY, size=12, space_after=8)

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Action", "Owner", "Timeline"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            _set_cell_shading(cell, "002960")
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = "Arial"
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE

        for slide in next_steps_slides:
            items = slide.get("action_items", slide.get("bullets", []))
            for item in items:
                row = table.add_row()
                if isinstance(item, dict) and "action" in item:
                    row.cells[0].text = item.get("action", "")
                    row.cells[1].text = item.get("owner", "TBD")
                    row.cells[2].text = item.get("timeline", "TBD")
                elif isinstance(item, dict):
                    row.cells[0].text = item.get("text", str(item))
                else:
                    row.cells[0].text = str(item)

                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(9)

        doc.add_paragraph()  # spacing

    # ── Research Summary (if available) ──
    if research_brief and isinstance(research_brief, dict):
        exec_summary = research_brief.get("executive_summary", "")
        if exec_summary:
            _add_styled_para(doc, "RESEARCH SUMMARY", bold=True, color=NAVY, size=12, space_after=8)
            _add_styled_para(doc, exec_summary, size=10, space_after=8)

        strongest = research_brief.get("strongest_evidence", [])
        if strongest:
            _add_styled_para(doc, "Key Evidence:", bold=True, color=DARK_GRAY, size=10, space_after=4)
            for ev in strongest[:5]:
                para = doc.add_paragraph(str(ev), style="List Bullet")
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            doc.add_paragraph()

    # ── Footer ──
    _add_styled_para(doc, "─" * 72, color=RGBColor(0xCC, 0xCC, 0xCC), size=6, space_after=4)
    _add_styled_para(doc, f"Generated by McKinsey Deck Builder | {date_str}",
                     color=MED_GRAY, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
